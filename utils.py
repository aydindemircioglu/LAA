import shutil
import os
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

import cv2
import numpy as np


def recreatePath (path, dry = False):
    print ("Recreating path ", path)
    if dry == True:
        return None

    try:
        shutil.rmtree (path)
    except:
        pass
    os.makedirs (path)
    pass


def removePath (path, dry):
    print ("Removing path ", path)
    if dry == True:
        return None

    try:
        shutil.rmtree (path)
    except:
        pass
    pass






def saveTable (diceTbl, fname = "XXX"):
    # save
    doc = docx.Document()

    styles = doc.styles
    style = styles.add_style('Bold', WD_STYLE_TYPE.PARAGRAPH)
    style.font.bold = True
    style.font.size = Pt(7)
    doc.styles['Normal'].font.size=Pt(7)

    doctbl = doc.add_table(diceTbl.shape[0]+1, diceTbl.shape[1]+1)

    # add the header rows.
    for j in range(diceTbl.shape[0]):
            #doctbl.cell(0,j).text = tasktbl.columns[j]
            text = diceTbl.index[j]
            #doctbl.cell(0,j).add_paragraph (text, style = 'Bold')
            doctbl.cell(j+1,0).paragraphs[0].text = text
            doctbl.cell(j+1,0).paragraphs[0].style = style

    # add the header rows.
    for j in range(diceTbl.shape[1]):
            #doctbl.cell(0,j).text = tasktbl.columns[j]
            text = diceTbl.columns[j]
            #doctbl.cell(0,j).add_paragraph (text, style = 'Bold')
            doctbl.cell(0,j+1).paragraphs[0].text = text
            doctbl.cell(0,j+1).paragraphs[0].style = style

    # add the rest of the data frame
    for j in range(diceTbl.shape[1]):
            colname = diceTbl.columns[j]
            for i in range(diceTbl.shape[0]):
                    style = 'Normal'
                    if  colname == "AUC" or colname == "Sensitivity" or colname == "Specificity":
                            text = str(round(diceTbl.values[i,j], 3))
                    else:
                            text = str(diceTbl.values[i,j])
                    doctbl.cell(i+1,j+1).paragraphs[0].text = text
                    doctbl.cell(i+1,j+1).paragraphs[0].style = style
                    #doctbl.cell(i+1,j).add_paragraph (text, style = style)
    doc.save("./paper/" + fname + ".docx")



def equalizeChannel (ch):
    maxCH = np.max(ch)
    ch = ch/maxCH*255
    ch = cv2.equalizeHist(ch.astype(np.uint8))
    ch = ch.astype(float)*ch
    return ch



def pseudoRGB (img, method = "clahe", visualize = False):
    if method not in ["clahe"]:
        exit ("Pseudo RGB method " + str(method) + " is unknown.")

    conversionFactor = 256
    if img.dtype == np.uint8:
        conversionFactor  = 1
        method = 'clahe'

    if method == "clahe":
        # quick hack for now
        #img[img >4000] = 0
        factor = 0.5
        clipfactor = 2
        baseFactor = 16.0
        spreadFactor = 2.0

        clahe = cv2.createCLAHE(clipLimit=baseFactor*spreadFactor*clipfactor, tileGridSize=(int(2*factor),int(2*factor)))
        red = clahe.apply(img)
        #red = (red/np.max(red)*255).astype('uint8')
        clahe = cv2.createCLAHE(clipLimit=baseFactor*1/spreadFactor*clipfactor, tileGridSize=(int(8*factor),int(8*factor)))
        blue = clahe.apply(img)
        #blue = (blue /np.max(blue)*255).astype('uint8')
        clahe = cv2.createCLAHE(clipLimit=baseFactor*clipfactor, tileGridSize=(int(4*factor),int(4*factor)))
        green = clahe.apply(img)


    if method == "clahe_new":
        # quick hack for now
        f = 1 # startfactor for tiling
        ef = 3 # multiplication factor
        c = 64 # start clip value

        # blue stays the same
        clahe = cv2.createCLAHE(clipLimit=c, tileGridSize=(f, f))
        #red = clahe.apply(img)
        red = img.copy()
        #red = (red/np.max(red)*255).astype('uint8')
        clahe = cv2.createCLAHE(clipLimit=c//ef//ef, tileGridSize=(ef*ef*f, ef*ef*f))
        blue = clahe.apply(img)
        #blue = (blue /np.max(blue)*255).astype('uint8')
        clahe = cv2.createCLAHE(clipLimit=c//ef, tileGridSize=(ef*f, ef*f))
        green = clahe.apply(img)

    # scale range
    scaleChannels = False
    if scaleChannels == True:
        red = (red-np.min(red))/(np.max(red)-np.min(red))
        green = (green-np.min(green))/(np.max(green)-np.min(green))
        blue = (blue-np.min(blue))/(np.max(blue)-np.min(blue))

    doEqualize = False
    if (doEqualize == True):
        green = equalizeChannel (green)
        blue = equalizeChannel (blue)
        red = equalizeChannel (red)
    img = cv2.merge((blue, green, red))

    if visualize == True:
        cv2.imshow('image512',img)
        cv2.waitKey(0)
        cv2.destroyAllWindows()
        for i in range (1,5):
            cv2.waitKey(1)

    return img


# some test
if 1 == 0:
    from PIL import Image
    import numpy as np
    import cv2
    f = "/data/uke/tr.pulmonalis/export/delimitations/normal/ac3a2925-52bfcef0-112c774a-a44effad-f1d16bd8.png"
    f = "/data/uke/tr.pulmonalis/export/delimitations/normal/e879767b-dbba3d2d-c243d6c0-85f5ab48-8aa29a9f.png"
    img = cv2.imread(f)[:,:,0]
    Image.fromarray(img)

    blue = img
    clahe = cv2.createCLAHE(clipLimit=16.0, tileGridSize=(4,4))
    blue = clahe.apply(img)

    clahe = cv2.createCLAHE(clipLimit=8.0, tileGridSize=(8,8))
    red = clahe.apply(img)
    Image.fromarray(red)

    clahe = cv2.createCLAHE(clipLimit=32.0, tileGridSize=(2,2))
    green = clahe.apply(img)
    Image.fromarray(green)

    rgbimg = cv2.merge((blue, green, red))
    Image.fromarray(rgbimg)

    import numpy as np
    Image.fromarray(pseudoRGB (img))
